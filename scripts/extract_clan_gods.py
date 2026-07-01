"""
Extract structured data from 'List of Clan Gods and Villages.docx'
into a normalized JSON file for the GIS map app.

Produces: data/gods_and_goddesses/clan_gods.json
"""
import json
import re
import docx
from collections import OrderedDict

DOC_PATH = "data/Extra Data/gods_and_goddesses/List of Clan Gods and Villages.docx"
OUT_PATH = "data/gods_and_goddesses/clan_gods.json"

doc = docx.Document(DOC_PATH)

# ── 1. Canonical name registry ──
# Maps every known variant → canonical ID
ALIAS_MAP = {}

def make_id(name):
    """Generate a canonical ID from a pen name."""
    if not name:
        return None
    name = name.strip()
    # Remove newlines
    name = name.replace('\n', ' ').replace('\r', ' ')
    # Clean: lowercase, keep only a-z, 0-9, spaces, and underscores
    n = name.lower()
    n = re.sub(r'[^a-z0-9_\s]', '', n)
    n = re.sub(r'\s+', '_', n.strip())
    return n

def register_pen(name, canonical_id=None, source_context=""):
    """Register a pen name variant and return canonical ID."""
    if not name or name in ('?', '-', '', '"', '“', '”'):
        return None
    name = name.strip().rstrip('*+')
    if not name:
        return None

    # If canonical_id given, use it; otherwise check if name is already a known alias
    if canonical_id is None:
        existing_cid = resolve_id(name)
        if existing_cid:
            return existing_cid
        cid = make_id(name)
    else:
        cid = canonical_id

    if cid not in ALIAS_MAP:
        ALIAS_MAP[cid] = {"primary": name, "aliases": OrderedDict()}

    # Store alias if different from primary
    existing = ALIAS_MAP[cid]
    normalized = make_id(name)
    if normalized != cid:
        existing["aliases"][name] = source_context

    return cid

def resolve_id(name):
    """Resolve a pen name to its canonical ID."""
    if not name or name in ('?', '-', '', '"', '“', '”'):
        return None
    name = name.strip().rstrip('*+')

    # Direct match
    cid = make_id(name)
    if cid in ALIAS_MAP:
        return cid

    # Try alias lookup
    for cid, info in ALIAS_MAP.items():
        if name in info["aliases"]:
            return cid
        # Check normalized aliases
        for alias_name in info["aliases"]:
            if make_id(alias_name) == make_id(name):
                return cid
        # Check primary
        if make_id(info["primary"]) == make_id(name):
            return cid
    return None

# ── 2. Pre-register all known pen names from tables and narratives ──

# From Table 0 (Kunjam)
for name, ctx in [
    ("Vedmo Moitor", "Table 0"),
    ("Urru Moitor", "Table 0"),
    ("Urru Dokri", "Table 0"),
    ("Punga Rungi", "Table 0"),
    # Hurra Mara/Huru Mara merged below via canonical_id
    ("Mai Sunga", "Table 0"),
    ("Desh Deva", "Table 0"),
    ("Munuk Deva", "Table 0"),
    ("Chichur Unga", "Table 0"),
    ("Hadma Raj", "Table 0"),
    ("Chichur Urra", "Table 0"),
    ("Jalsingo", "Table 0"),
]:
    register_pen(name, source_context=ctx)

# Register Huru Mara = Hurra Mara (canonical: huru_mara)
huru_mara_id = make_id("Huru Mara")
register_pen("Huru Mara", canonical_id=huru_mara_id, source_context="Narrative")
register_pen("Hurra Mara", canonical_id=huru_mara_id, source_context="Table 0")
register_pen("Urru Mara", canonical_id=huru_mara_id, source_context="Table 1")
register_pen("Hura Mara", canonical_id=huru_mara_id, source_context="Paragraph 11")
register_pen("Urru/Huru Mara", canonical_id=huru_mara_id, source_context="Table 1 Row 8")

# Urru Moitor is a different pen (Kunjam, Kuhrami)
urru_moitor_id = make_id("Urru Moitor")
register_pen("Urru Moitor", canonical_id=urru_moitor_id, source_context="Table 0")

# From Table 1 (Markami)
# Note: names with variant spellings merged later via canonical_id are omitted here
for name, ctx in [
    ("Iram Raj", "Table 1"),
    ("Hunga Moitor", "Table 1"),
    ("Gujje Dokri", "Table 1"),
    ("Bhimaraj", "Table 1"),
    ("Dol Mutte", "Narrative"),
    ("Dayur Mutte", "Narrative"),
    ("Bhogam Mutte", "Narrative"),
    ("Urru Ponde", "Table 1"),
    ("Biriya Bhima", "Table 1"),
    ("Mudde Dokri", "Table 1"),
    ("Bomul Ungal", "Table 1"),
    ("Godar Bhima", "Table 1"),
    ("Chaikut Bhima", "Table 1"),
    # Chaikud Bhima registered via canonical merge below
    ("Bhime", "Table 1"),
    ("Inge Dokri", "Table 1"),
    ("Lug Unga", "Table 1"),
    ("Godar Hunga", "Table 1"),
    ("Daro Moitor", "Table 1"),
    # Darrem Modka registered via canonical merge below
    ("Raibandal", "Table 1"),
    ("Akaluru Dokri", "Table 1"),
    ("Kohla Kosu", "Table 1"),
    ("Peda Hadma", "Table 1"),
    ("Chinna Hadma", "Table 1"),
    ("Murde Muyal", "Table 1"),
    ("Hunga Bhimal", "Table 1"),
    ("Pind Hadmal", "Table 1"),
    # Bhum Iriya/Bhim Iriya registered below with canonical merge
    ("Pal Hadmal", "Table 1"),
]:
    register_pen(name, source_context=ctx)

register_pen("Iram Raj", canonical_id=make_id("Iram Raj"), source_context="Table 1")
register_pen("Irma Raj", canonical_id=make_id("Iram Raj"), source_context="Table 1 (variant)")

bhimaraj_id = make_id("Bhimaraj")
register_pen("Bhimaraj", canonical_id=bhimaraj_id, source_context="Table 1")
register_pen("Bhimraj", canonical_id=bhimaraj_id, source_context="Narrative")

godar_bhima_id = make_id("Godar Bhima")
register_pen("Godar Bhima", canonical_id=godar_bhima_id, source_context="Table 1")
register_pen("Chaikut Bhima", canonical_id=make_id("Chaikut Bhima"), source_context="Table 1")
register_pen("Chaikud Bhima", canonical_id=make_id("Chaikut Bhima"), source_context="Narrative")

daro_moitor_id = make_id("Daro Moitor")
register_pen("Daro Moitor", canonical_id=daro_moitor_id, source_context="Table 1")
register_pen("Darrem Modka", canonical_id=daro_moitor_id, source_context="Table 1")
register_pen("Daro Modka", canonical_id=daro_moitor_id, source_context="Paragraph 11")

bhum_iriya_id = make_id("Bhum Iriyal")
register_pen("Bhum Iriyal", canonical_id=bhum_iriya_id, source_context="Table 1")
register_pen("Bhum Iriya", canonical_id=bhum_iriya_id, source_context="Table 2")
register_pen("Bhim Iriya", canonical_id=bhum_iriya_id, source_context="Paragraph 11")
register_pen("Bhum Iriyal", canonical_id=bhum_iriya_id, source_context="Table 1 Row 21")

# Mawe Lungo = Mawe Kungo
mawe_lungo_id = make_id("Mawe Lungo")
register_pen("Mawe Lungo", canonical_id=mawe_lungo_id, source_context="Narrative")
register_pen("Mawe Kungo", canonical_id=mawe_lungo_id, source_context="Table 1 Row 9")

# Katta Bodke = Katte Bodke
katta_bodke_id = make_id("Katta Bodke")
register_pen("Katta Bodke", canonical_id=katta_bodke_id, source_context="Narrative")
register_pen("Katte Bodke", canonical_id=katta_bodke_id, source_context="Table 1 Row 11")

# Vange Dokri = Vunge Dokri
vange_dokri_id = make_id("Vange Dokri")
register_pen("Vange Dokri", canonical_id=vange_dokri_id, source_context="Table 3")
register_pen("Vunge Dokri", canonical_id=vange_dokri_id, source_context="Table 3 (variant)")

# Urru Mara = Urru = Huru Mara (already merged above)
# Punga Rungi relationship: ensure Mai Sunga ghar jamai resolves correctly

# ── Cleanup: remove stale ALIAS_MAP entries superseded by canonical merges ──
stale_ids = []
for cid, info in list(ALIAS_MAP.items()):
    for other_cid, other_info in ALIAS_MAP.items():
        if cid == other_cid:
            continue
        if info["primary"] in other_info["aliases"]:
            # cid's primary is an alias of other_cid → cid is stale
            stale_ids.append(cid)
            break
        # Also check if this cid's primary resolves to a different canonical via make_id
        if make_id(info["primary"]) != cid and make_id(info["primary"]) in ALIAS_MAP:
            stale_ids.append(cid)
            break
for sid in set(stale_ids):
    del ALIAS_MAP[sid]

# From Table 2 (Oyam)
for name, ctx in [
    ("Muddaraj", "Table 2"),
    ("Vidi Iriyal Panda Hadma", "Table 2"),
    ("Gaddi Kama", "Table 2"),
    ("Murke Nango", "Table 2"),
    ("Kohli Dokri", "Table 2"),
    ("Gal Dullo", "Table 2"),
    ("Gal Bomda", "Table 2"),
    ("Hingal Devo", "Table 2"),
    ("Chenna Kama", "Table 2"),
    ("Ghanta Kama", "Table 2"),
    ("Ukud Kama", "Table 2"),
    ("Bade Vidi Iriyal", "Table 2"),
    ("Hadma Iriyal", "Table 2"),
    ("Bande Boyo", "Table 2"),
]:
    register_pen(name, source_context=ctx)

gaddi_kama_id = make_id("Gaddi Kama")
register_pen("Gaddi Kama", canonical_id=gaddi_kama_id, source_context="Table 2")
register_pen("Ghadi Kama", canonical_id=gaddi_kama_id, source_context="Narrative (shrine)")
register_pen("Gadye Kama", canonical_id=gaddi_kama_id, source_context="Ballu Bhavani spelling")
register_pen("Gadi Kama", canonical_id=gaddi_kama_id, source_context="Paragraph 53")
register_pen("Gadye/Gaddi Kama", canonical_id=gaddi_kama_id, source_context="Paragraph 53")
register_pen("Ghadi/Gadi/Gadye Kama", canonical_id=gaddi_kama_id, source_context="Paragraph 59")

# From Table 3 (Madvi contd)
# Note: variants merged below via canonical_id are omitted here
for name, ctx in [
    ("Vange Dokri", "Table 3"),
    ("Nanga Bhima", "Table 3"),
    ("Barra Bujja", "Table 3"),
    ("Morka Moitor", "Table 3"),
    ("Markaraj", "Table 3"),
    ("Hinge Dokri", "Table 3"),
    ("Lingal Denga", "Narrative"),
    ("Andalkosa", "Table 3"),
]:
    register_pen(name, source_context=ctx)

markaraj_id = make_id("Markaraj")
register_pen("Markaraj", canonical_id=markaraj_id, source_context="Table 3")
register_pen("Madkaraj", canonical_id=markaraj_id, source_context="Paragraph 53")
register_pen("Markaraj**", canonical_id=markaraj_id, source_context="Table 3")

ural_gunda_id = make_id("Ural Gunda")
register_pen("Ural Gunda", canonical_id=ural_gunda_id, source_context="Narrative")
register_pen("Uraal Gundal", canonical_id=ural_gunda_id, source_context="Table 3")

vange_dokri_id = make_id("Vange Dokri")
register_pen("Vange Dokri", canonical_id=vange_dokri_id, source_context="Table 3")
register_pen("Vunge Dokri", canonical_id=vange_dokri_id, source_context="Table 3 (variant)")

andalkosa_id = make_id("Andalkosa")
register_pen("Andalkosa", canonical_id=andalkosa_id, source_context="Table 3")
register_pen("Andal Kosa", canonical_id=andalkosa_id, source_context="Paragraph 53")

# Mawe Lungo (mentioned in narrative)
register_pen("Mawe Lungo", source_context="Narrative")
register_pen("Mawe Kungo", source_context="Table 1 Row 9")

# Katta Bodke
register_pen("Katta Bodke", source_context="Narrative")
register_pen("Katte Bodke", source_context="Table 1 row 11")

# Nandraj (mountain/personified figure)
register_pen("Nandraj", source_context="Narrative")
register_pen("Mawli", source_context="Narrative")

# ══════════════════════════════════════════════
# Helper functions (must be defined before use)
# ══════════════════════════════════════════════

def infer_gender(name):
    if not name:
        return "unknown"
    female_indicators = ["dokri", "dokrar", "moyo", "lungo", "rungi", "ponde", "nango", "bhime", "bodke", "mutte"]
    for fi in female_indicators:
        if fi in name.lower():
            return "female"
    return "male"

def get_pen(name, clan_id=None, phratry_id=None, source_context=""):
    cid = resolve_id(name)
    if not cid:
        cid = make_id(name)
        register_pen(name, canonical_id=cid, source_context=source_context or "auto")
    if cid not in pens:
        primary = ALIAS_MAP.get(cid, {}).get("primary", name)
        pens[cid] = {
            "id": cid,
            "name": primary,
            "aliases": list(ALIAS_MAP.get(cid, {}).get("aliases", {}).keys()),
            "clan_id": clan_id,
            "phratry_id": phratry_id,
            "gender": infer_gender(primary),
            "gudi_village": None,
            "palli": [],
            "perma": "",
            "karsad": "",
            "notes": "",
            "type": "main"
        }
    return cid

def add_relationship(rel_type, from_name, to_name, details="", source=""):
    # If already resolved to an ID, use directly; otherwise resolve
    from_id = from_name if from_name and from_name in pens else resolve_id(from_name)
    to_id = to_name if to_name and to_name in pens else resolve_id(to_name)
    if not from_id or not to_id:
        return
    rel = {
        "type": rel_type,
        "from_pen_id": from_id,
        "to_pen_id": to_id,
        "details": details,
        "source": source
    }
    if rel not in relationships:
        relationships.append(rel)

def parse_marriage_info(pid, text, source=""):
    """Parse marriage info from a relations field."""
    pen_name = pens[pid]["name"]
    # Extract the spouse name: up to 2 words, stopping at descriptor words
    m = re.search(r'(?:Married to|married to|is married to)\s+([A-Za-z]+(?:\s+[A-Za-z]+){0,1})(?:\s+(?:pen|husband|wife|in|of|d/o|s/o|son|daughter|kunjam)|\s*\(|\s*,)', text)
    if not m:
        # Try simpler: just match first 1-2 words after "married to"
        m = re.search(r'(?:Married to|married to|is married to)\s+([A-Za-z]+(?:\s+[A-Za-z]+){0,1})', text)
    if m:
        spouse_text = m.group(1).strip().rstrip('.')
        spouse_name = clean_pen_name(spouse_text)
        if spouse_name and is_likely_name(spouse_name):
            sid = get_pen(spouse_name, source_context=source)
            add_relationship("marriage", pid, sid, details=f"{pen_name} married to {spouse_name}", source=source)

def parse_sibling_info(pid, text, source=""):
    """Parse sibling info."""
    pen_name = pens[pid]["name"]
    m = re.search(r'(?:Younger|Elder|younger|elder)\s+brother\s+of\s+([A-Za-z]+(?:\s+[A-Za-z]+){0,3})', text)
    if m:
        sibling_name = m.group(1).strip().rstrip('.')
        sibling_name = clean_pen_name(sibling_name)
        if sibling_name and is_likely_name(sibling_name):
            sid = resolve_id(sibling_name)
            if not sid:
                sid = get_pen(sibling_name, source_context=source)
            is_younger = 'younger' in m.group(0).lower()
            if is_younger:
                add_relationship("sibling", sid, pid, details="younger brother", source=source)
            else:
                add_relationship("sibling", sid, pid, details="elder brother", source=source)

def clean_village_name(raw):
    """Clean a village name (handle newlines, question marks etc.)."""
    if not raw:
        return ""
    n = raw.strip()
    n = n.replace('\n', ' ').replace('\r', ' ')
    n = re.sub(r'\s+', ' ', n)
    n = n.strip().rstrip('?')
    return n.strip()

def is_likely_name(text):
    """Check if text looks like a proper name (all words start with uppercase, no stop words)."""
    if not text:
        return False
    stop_words = {'a', 'an', 'the', 'is', 'are', 'was', 'were', 'to', 'of', 'in', 'for', 'on',
                  'and', 'or', 'but', 'his', 'her', 'their', 'this', 'that', 'with', 'from',
                  'by', 'at', 'has', 'had', 'have', 'been', 'all', 'also', 'now', 'not',
                  'younger', 'elder', 'still', 'married', 'called', 'known'}
    words = text.split()
    if len(words) > 4:
        return False
    for w in words:
        w_clean = w.strip('.,;:!?()[]{}')
        if not w_clean:
            continue
        # Must start with uppercase (proper name)
        if w_clean[0].islower():
            return False
        if w_clean.lower() in stop_words:
            return False
    return True

def parse_parent_child_from_name(text, clan_id, phratry_id, source=""):
    """Parse 's/o X' or 'son of X' patterns in pen name fields.
    Only extracts when child name is the immediate 1-4 words before the marker
    AND looks like a proper name.
    """
    m = re.search(r'(?:s/o|S/o|son of|Son of|d/o|D/o|daughter of|Daughter of)\s+([A-Za-z]+(?:\s+[A-Za-z]+){0,3})', text)
    if m:
        parent_name = m.group(1).strip().rstrip('.')
        parent_name = clean_pen_name(parent_name)
        # Extract the last 1-4 words before the marker as child name
        before = text[:m.start()].strip().rstrip(',').strip()
        before_words = before.split()
        if len(before_words) > 4:
            before_words = before_words[-4:]
        child_name = clean_pen_name(' '.join(before_words))
        if child_name and parent_name and parent_name != '?' and is_likely_name(child_name):
            pid = get_pen(parent_name, source_context=source)
            cid = get_pen(child_name, clan_id=clan_id, phratry_id=phratry_id)
            add_relationship("parent", pid, cid, details=f"{parent_name} is parent of {child_name}", source=source)

def clean_pen_name(raw):
    """Extract just the pen name from text that may contain notes, relationships, etc."""
    if not raw:
        return None
    n = raw.strip()
    # Replace newlines with space
    n = n.replace('\n', ' ').replace('\r', ' ')
    # Collapse multiple spaces
    n = re.sub(r'\s+', ' ', n)
    # Remove content in parentheses (including unmatched opening paren)
    n = re.sub(r'\([^)]*\)', '', n)
    n = re.sub(r'\([^)]*$', '', n)
    # Remove content after semicolons
    n = re.sub(r';.*$', '', n)
    # Remove asterisks, quotes, and other markers
    n = n.replace('*', '').replace('+', '').replace('"', '').replace('\u201c', '').replace('\u201d', '')
    # Remove uncertainty/descriptive phrases at start
    n = re.sub(r'^(Either|Either a|Either an|Unknown|The)\s+', '', n)
    n = re.sub(r'\s+or a\s+.*$', '', n)
    # Also remove anything that starts with "(and" or similar notes at end
    n = re.sub(r'\s*\(?\s*and\s+[A-Za-z]+\s*\)?\s*$', '', n)
    # Strip ", s/o ...", ", d/o ...", ", son of ..." etc.
    n = re.sub(r',\s*(?:s/o|S/o|son of|d/o|D/o|daughter of)\s+.*$', '', n)
    n = re.sub(r'\s+and\s*$', '', n)
    # Handle "/" aliases: take the first variant that forms a complete name
    # Only split if what comes after the "/" is an entire alternate name,
    # not just the middle word of a multi-word name (e.g. "Iram/Irma Raj" ≠ "Iram")
    if '/' in n:
        parts = [p.strip() for p in n.split('/')]
        if len(parts) >= 2:
            # Check if taking first part creates a 1-word name while
            # a later part has more words (means "/" is within first name)
            # e.g. parts=["Iram","Irma Raj"] → keep "Iram Raj" by joining
            if len(parts[0].split()) == 1 and any(len(p.split()) > 1 for p in parts[1:]):
                # First part is just a first word variant; take the last word from a multi-word variant
                for p in parts[1:]:
                    if len(p.split()) > 1:
                        n = parts[0] + ' ' + ' '.join(p.split()[1:])
                        break
            else:
                n = parts[0]
    n = n.strip()
    # Validate: must be 1-4 words, not start with lowercase relational words
    words = n.split()
    skip_words = {'younger', 'elder', 'still', 'and', 'or', 'the', 'a', 'an', 'his', 'her', 'they', 'their', 'this', 'that'}
    if len(words) == 0 or len(words) > 4 or words[0].lower() in skip_words:
        return None
    return n

def normalize_clan_name(cname):
    mapping = {
        "markami": "markami",
        "kalmu/karma": "kalmu",
        "kalmu": "kalmu",
        "karma": "kalmu",
        "bhogam/chote telam": "bhogam",
        "bhogam": "bhogam",
        "chote telam": "telam_chote",
        "chote telam/bhogam": "telam_chote",
        "bade telam": "telam_bade",
        "telam": "telam_bade",
        "icham": "icham",
        "punem": "punem",
        "kadiyam": "kadiyam",
        "midiyam": "midiyam",
        "undam": "undam",
        "tati": "tati",
        "kadti": "kadti",
        "rengo": "rengo",
        "kunjam": "kunjam",
        "oyam": "oyam",
        "barse": "barse",
        "madvi": "madvi",
        "hemla": "hemla",
        "tamo": "tamo",
        "padami": "padami",
        "ujji/dodi": "ujji_dodi",
        "kawasi": "kawasi",
    }
    return mapping.get(cname.strip().lower(), cname.strip().lower().replace("/", "_").replace(" ", "_"))

def parse_implicit_pens(text, phratry_id, clan_id, source=""):
    """Parse pen names from relational text (continuation rows)."""
    parts = re.split(r'[;,]\s*', text)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r'^(s/o|d/o|son of|daughter of|married to|ghar jamai)', part, re.IGNORECASE):
            continue
        name_match = re.match(r'^([A-Za-z\s]+?)\s+(?:s/o|d/o|son of|daughter of|married to|ghar jamai|is\s+)', part, re.IGNORECASE)
        if name_match:
            pn = name_match.group(1).strip().rstrip(',')
            if pn:
                pid = get_pen(pn, clan_id=clan_id, phratry_id=phratry_id)
                pens[pid]["type"] = "subordinate"
                parse_parent_child_from_name(part, clan_id, phratry_id, source=source)
                if "ghar jamai" in part.lower():
                    mm = re.search(r'married to\s+([A-Za-z\s]+)', part)
                    if mm:
                        spouse_name = mm.group(1).strip().rstrip('.')
                        add_relationship("ghar_jamai", pn, spouse_name,
                                         details=f"{pn} is ghar jamai of {spouse_name}", source=source)


# ── 3. Extract structured data from tables ──

pens = {}
villages = []
relationships = []
clans = {}
phratries = OrderedDict()

# Phratry definitions
phratries["kuhrami_kadiari"] = {
    "id": "kuhrami_kadiari",
    "name": "Kuhrami/Kadiari",
    "source": "Grigson, p. 301",
    "clan_ids": ["kunjam"]
}
phratries["markami"] = {
    "id": "markami",
    "name": "Markami Kutumb",
    "source": "",
    "clan_ids": ["markami", "kalmu", "bhogam", "telam_bade", "telam_chote",
                  "icham", "punem", "kadiyam", "midiyam", "undam", "tati", "kadti", "rengo"]
}
phratries["madvi"] = {
    "id": "madvi",
    "name": "Madvi",
    "source": "",
    "clan_ids": ["oyam", "barse", "madvi", "hemla", "tamo", "padami", "ujji_dodi"]
}
phratries["kawasi"] = {
    "id": "kawasi",
    "name": "Kawasi",
    "source": "",
    "clan_ids": ["kawasi"]
}

# ── 3a. Table 0: Kunjam Clan Gods (Kuhrami/Kadiari) ──
table0 = doc.tables[0]
t0_phratry = "kuhrami_kadiari"
t0_clan = "kunjam"

# Register kunjam clan (referenced by Tables 0-4 but never defined)
if "kunjam" not in clans:
    clans["kunjam"] = {"id": "kunjam", "name": "Kunjam", "phratry_id": "kuhrami_kadiari", "notes": ""}

t0_rows = []
for r_idx in range(1, len(table0.rows)):
    row = table0.rows[r_idx]
    cells = [cell.text.strip() for cell in row.cells]
    t0_rows.append({
        "sno": cells[0] if len(cells) > 0 else "",
        "village": cells[1] if len(cells) > 1 else "",
        "main_pen": cells[2] if len(cells) > 2 else "",
        "relations": cells[3] if len(cells) > 3 else "",
        "subordinate": cells[4] if len(cells) > 4 else "",
        "remarks": cells[5] if len(cells) > 5 else ""
    })

for r in t0_rows:
    vname = clean_village_name(r["village"])
    if not vname:
        continue

    raw_main_pen = r["main_pen"].strip()
    main_pen_name = clean_pen_name(raw_main_pen)
    sub_info = r["subordinate"].strip()
    rel_info = r["relations"].strip()
    remarks = r["remarks"].strip()

    villagers_entry = {
        "name": vname,
        "aliases": [],
        "phratry_id": t0_phratry,
        "clan_id": t0_clan,
        "main_pen_id": None,
        "subordinate_pen_ids": [],
        "notes": remarks,
        "population_info": ""
    }

    # Handle main pen
    if main_pen_name:
        if main_pen_name.startswith("Son of") or "Kunjam pen" in main_pen_name or "Unknown" in main_pen_name:
            pen_note = f"Pen: {raw_main_pen}"
            if remarks:
                pen_note += f". {remarks}"
            villagers_entry["notes"] = pen_note
        else:
            pid = get_pen(main_pen_name, clan_id=t0_clan, phratry_id=t0_phratry)
            villagers_entry["main_pen_id"] = pid
            pens[pid]["gudi_village"] = vname

            if rel_info:
                parse_marriage_info(pid, rel_info, source=f"Table 0: {vname}")
                parse_sibling_info(pid, rel_info, source=f"Table 0: {vname}")

    # Handle subordinate pens
    if sub_info:
        # Split by semicolons first (these separate different pen entries)
        sub_groups = re.split(r';\s*', sub_info)
        for group in sub_groups:
            group = group.strip().rstrip('.')
            if not group:
                continue
            # Handle "and" within a group (separates multiple pen names)
            sub_names = re.split(r'\s+and\s+', group)
            for sp in sub_names:
                sp = sp.strip()
                if not sp:
                    continue
                # Extract just the pen name (before any comma or parenthetical describing relationship)
                sp_name = clean_pen_name(sp)
                # Also extract name before first comma if clean_pen_name didn't catch everything
                if not sp_name:
                    continue

                ghar_jamai = "ghar jamai" in sp.lower()
                bachelor = "bachelor" in sp.lower()
                younger_of = ""
                ym = re.search(r'younger\s+brother\s+of\s+([^,;]+)', sp, re.IGNORECASE)
                if ym:
                    younger_of = clean_pen_name(ym.group(1).strip())

                # Final check: skip fragments that aren't real names
                skip_words = ["younger", "bachelor", "still a", "younger to", "elder"]
                if any(sp_name.lower().startswith(w) for w in skip_words):
                    continue
                if len(sp_name.split()) > 4:
                    continue  # Too long to be a name

                spid = get_pen(sp_name, clan_id=t0_clan, phratry_id=t0_phratry)
                pens[spid]["type"] = "subordinate"
                pens[spid]["gudi_village"] = vname
                if bachelor:
                    pens[spid]["notes"] = "Bachelor"
                if ghar_jamai and villagers_entry["main_pen_id"]:
                    add_relationship("ghar_jamai", sp_name, main_pen_name,
                                     details=f"{sp_name} is ghar jamai", source=f"Table 0: {vname}")
                if younger_of:
                    younger_id = resolve_id(younger_of)
                    if younger_id:
                        add_relationship("sibling", younger_id, sp_name,
                                         details="younger brother", source=f"Table 0: {vname}")

                villagers_entry["subordinate_pen_ids"].append(spid)

    # Parse parent-child from Pen name field
    if main_pen_name:
        parse_parent_child_from_name(raw_main_pen, t0_clan, t0_phratry, source=f"Table 0: {vname}")

    villages.append(villagers_entry)

# ── 3b. Table 1: Markami Kutumb ──
table1 = doc.tables[1]
t1_phratry = "markami"

t1_rows = []
for r_idx in range(1, len(table1.rows)):
    row = table1.rows[r_idx]
    cells = [cell.text.strip() for cell in row.cells]
    t1_rows.append({
        "sno": cells[0] if len(cells) > 0 else "",
        "village": cells[1] if len(cells) > 1 else "",
        "clan": cells[2] if len(cells) > 2 else "",
        "pen": cells[3] if len(cells) > 3 else "",
        "relations": cells[4] if len(cells) > 4 else "",
        "remarks": cells[5] if len(cells) > 5 else ""
    })

for r in t1_rows:
    vname = clean_village_name(r["village"])
    if not vname:
        # Handle continuation rows (rows 9, 10 where village is empty but relationships continue)
        rel_info = r["relations"].strip()
        if rel_info:
            # These are additional relationship info for the previous village
            if villages:
                prev_v = villages[-1]
                # Parse "Lug Unga, s/o Huru Mara and Mawe Kungo"
                parse_implicit_pens(rel_info, t1_phratry, r["clan"] or prev_v.get("clan_id", ""),
                                    source=f"Table 1 cont: {prev_v['name']}")
        continue

    clan_raw = r["clan"].strip()
    clan_id = normalize_clan_name(clan_raw) if clan_raw else None
    raw_pen_name = r["pen"].strip()
    pen_name = clean_pen_name(raw_pen_name) or ""
    rel_info = r["relations"].strip()
    remarks = r["remarks"].strip()

    # Register clan
    if clan_id and clan_id not in clans:
        clans[clan_id] = {"id": clan_id, "name": clan_raw, "phratry_id": t1_phratry, "notes": ""}

    villagers_entry = {
        "name": vname,
        "aliases": [],
        "phratry_id": t1_phratry,
        "clan_id": clan_id,
        "main_pen_id": None,
        "subordinate_pen_ids": [],
        "notes": remarks,
        "population_info": ""
    }

    if pen_name and pen_name != '?':
        # Handle multiple pens (comma or "and" separated)
        if " and " in pen_name or "," in pen_name:
            pen_names = [p.strip().rstrip('*+()') for p in re.split(r'\s+and\s+|,\s*', pen_name) if p.strip()]
        elif "/" in pen_name and not any(p in pen_name for p in ["s/o"]):
            pen_names = [pen_name]
        else:
            pen_names = [pen_name]

        for pn in pen_names:
            if not pn:
                continue
            pid = get_pen(pn, clan_id=clan_id, phratry_id=t1_phratry)
            pens[pid]["gudi_village"] = vname

            if villagers_entry["main_pen_id"] is None:
                villagers_entry["main_pen_id"] = pid
            else:
                villagers_entry["subordinate_pen_ids"].append(pid)
                pens[pid]["type"] = "subordinate"

            if rel_info:
                parse_marriage_info(pid, rel_info, source=f"Table 1: {vname}")
                parse_sibling_info(pid, rel_info, source=f"Table 1: {vname}")
                parse_parent_child_from_name(rel_info, clan_id, t1_phratry, source=f"Table 1: {vname}")

            # Parse phalli from remarks
            if remarks:
                # Try to find actual palli content first (e.g., "palli covers X, Y, Z")
                pm = re.search(r"palli\s+(?:covers|includes|extends to)\s+([A-Za-z\s,–-]+)", remarks, re.IGNORECASE)
                if pm:
                    palli_villages = [v.strip() for v in re.split(r'[,–-]\s*', pm.group(1).strip()) if v.strip()]
                    pens[pid]["palli"].extend(palli_villages)
                else:
                    # Fallback: try to find "Name's palli" pattern
                    pm = re.search(r"([A-Za-z\s,]+)'?s?\s+palli\s+", remarks)
                    if pm:
                        palli_text = pm.group(1).strip()
                        if palli_text and len(palli_text) > 2:
                            palli_villages = [v.strip() for v in palli_text.split(",")]
                            pens[pid]["palli"].extend(palli_villages)

            # Parse perma
            permam = re.search(r"perma\s+is\s+([A-Za-z\s]+)", remarks, re.IGNORECASE)
            if permam:
                pens[pid]["perma"] = permam.group(1).strip()

    # Parse additional relations for empty-pen rows
    if not pen_name and rel_info:
        parse_implicit_pens(rel_info, t1_phratry, clan_id, source=f"Table 1: {vname}")

    villages.append(villagers_entry)

# ── 3c. Table 2: Oyam Clan Gods (Madvi) ──
table2 = doc.tables[2]
t2_phratry = "madvi"
t2_clan = "oyam"

t2_rows = []
for r_idx in range(1, len(table2.rows)):
    row = table2.rows[r_idx]
    cells = [cell.text.strip() for cell in row.cells]
    t2_rows.append({
        "sno": cells[0] if len(cells) > 0 else "",
        "village": cells[1] if len(cells) > 1 else "",
        "clan": cells[2] if len(cells) > 2 else "",
        "pen": cells[3] if len(cells) > 3 else "",
        "relations": cells[4] if len(cells) > 4 else "",
        "remarks": cells[5] if len(cells) > 5 else ""
    })

for r in t2_rows:
    vname = clean_village_name(r["village"])
    if not vname:
        continue
    clan_raw = r["clan"].strip()
    clan_id = normalize_clan_name(clan_raw) if clan_raw else t2_clan
    pen_name = clean_pen_name(r["pen"].strip()) or ""
    rel_info = r["relations"].strip()
    remarks = r["remarks"].strip()

    if clan_id not in clans:
        clans[clan_id] = {"id": clan_id, "name": clan_raw, "phratry_id": t2_phratry, "notes": ""}

    villagers_entry = {
        "name": vname,
        "aliases": [],
        "phratry_id": t2_phratry,
        "clan_id": clan_id,
        "main_pen_id": None,
        "subordinate_pen_ids": [],
        "notes": remarks,
        "population_info": ""
    }

    if pen_name and pen_name != '?':
        if " and " in pen_name:
            pen_names = re.split(r'\s+and\s+', pen_name)
        else:
            pen_names = [pen_name]

        for pn in pen_names:
            pn = pn.strip().rstrip('*+()')
            if not pn:
                continue
            pid = get_pen(pn, clan_id=clan_id, phratry_id=t2_phratry)
            pens[pid]["gudi_village"] = vname

            if villagers_entry["main_pen_id"] is None:
                villagers_entry["main_pen_id"] = pid
            else:
                villagers_entry["subordinate_pen_ids"].append(pid)
                pens[pid]["type"] = "subordinate"

            if rel_info:
                parse_marriage_info(pid, rel_info, source=f"Table 2: {vname}")
                parse_sibling_info(pid, rel_info, source=f"Table 2: {vname}")
                parse_parent_child_from_name(rel_info, clan_id, t2_phratry, source=f"Table 2: {vname}")

            if remarks:
                pm = re.search(r"phalli\s+(extends to|covers|includes)\s+([A-Za-z\s,–-]+)", remarks, re.IGNORECASE)
                if pm:
                    palli_text = pm.group(2)
                    palli_villages = [v.strip() for v in re.split(r'[,–-]\s*', palli_text) if v.strip()]
                    pens[pid]["palli"].extend(palli_villages)

    villages.append(villagers_entry)


# ── 3d. Table 3: Madvi contd + other clans ──
table3 = doc.tables[3]

t3_rows = []
for r_idx in range(1, len(table3.rows)):
    row = table3.rows[r_idx]
    cells = [cell.text.strip() for cell in row.cells]
    t3_rows.append({
        "sno": cells[0] if len(cells) > 0 else "",
        "village": cells[1] if len(cells) > 1 else "",
        "clan": cells[2] if len(cells) > 2 else "",
        "pen": cells[3] if len(cells) > 3 else "",
        "relations": cells[4] if len(cells) > 4 else "",
        "remarks": cells[5] if len(cells) > 5 else ""
    })

t3_default_phratry = "madvi"  # Most entries here are Madvi-related

for r in t3_rows:
    vname = clean_village_name(r["village"])
    if not vname:
        continue
    clan_raw = r["clan"].strip()
    clan_id = normalize_clan_name(clan_raw) if clan_raw else None
    pen_name = clean_pen_name(r["pen"].strip()) or ""
    rel_info = r["relations"].strip()
    remarks = r["remarks"].strip()

    # Determine phratry: all Table 3 entries are under Madvi phratry
    assigned_phratry = t3_default_phratry

    if clan_id and clan_id not in clans:
        clans[clan_id] = {"id": clan_id, "name": clan_raw, "phratry_id": assigned_phratry, "notes": ""}

    villagers_entry = {
        "name": vname,
        "aliases": [],
        "phratry_id": assigned_phratry,
        "clan_id": clan_id,
        "main_pen_id": None,
        "subordinate_pen_ids": [],
        "notes": remarks,
        "population_info": ""
    }

    if pen_name and pen_name != '?':
        pen_names = [pen_name]
        for pn in pen_names:
            pn = pn.strip().rstrip('*+()')
            if not pn:
                continue
            pid = get_pen(pn, clan_id=clan_id, phratry_id=assigned_phratry)
            pens[pid]["gudi_village"] = vname

            if villagers_entry["main_pen_id"] is None:
                villagers_entry["main_pen_id"] = pid
            else:
                villagers_entry["subordinate_pen_ids"].append(pid)
                pens[pid]["type"] = "subordinate"

            if rel_info:
                parse_marriage_info(pid, rel_info, source=f"Table 3: {vname}")
                parse_sibling_info(pid, rel_info, source=f"Table 3: {vname}")

            if remarks:
                pm = re.search(r"palli\s+(covers|includes)\s+([A-Za-z\s,]+)", remarks, re.IGNORECASE)
                if pm:
                    palli_text = pm.group(2)
                    palli_villages = [v.strip() for v in palli_text.split(",") if v.strip()]
                    pens[pid]["palli"].extend(palli_villages)

    villages.append(villagers_entry)


# ── 3e. Table 4: Kawasi ──
table4 = doc.tables[4]
t4_phratry = "kawasi"

for r_idx in range(1, len(table4.rows)):
    row = table4.rows[r_idx]
    cells = [cell.text.strip() for cell in row.cells]
    vname = clean_village_name(cells[1] if len(cells) > 1 else "")
    clan_raw = cells[2] if len(cells) > 2 else ""
    pen_name = cells[3] if len(cells) > 3 else ""

    if not vname:
        continue
    clan_id = normalize_clan_name(clan_raw) if clan_raw else "kawasi"
    if clan_id not in clans:
        clans[clan_id] = {"id": clan_id, "name": clan_raw, "phratry_id": t4_phratry, "notes": ""}

    villagers_entry = {
        "name": vname,
        "aliases": [],
        "phratry_id": t4_phratry,
        "clan_id": clan_id,
        "main_pen_id": None,
        "subordinate_pen_ids": [],
        "notes": "",
        "population_info": ""
    }

    if pen_name and pen_name not in ('?', '-'):
        pid = get_pen(pen_name, clan_id=clan_id, phratry_id=t4_phratry)
        pens[pid]["gudi_village"] = vname
        villagers_entry["main_pen_id"] = pid

    villages.append(villagers_entry)


# ── 3f. Ensure all registered pens exist in pens dict ──
for cid, info in ALIAS_MAP.items():
    if cid not in pens:
        pens[cid] = {
            "id": cid,
            "name": info["primary"],
            "aliases": list(info["aliases"].keys()),
            "clan_id": None,
            "phratry_id": None,
            "gender": infer_gender(info["primary"]),
            "gudi_village": None,
            "palli": [],
            "perma": "",
            "karsad": "",
            "notes": "Referenced in narratives/tables but no village assigned",
            "type": "main"
        }

# ── 4. Extract facts from narratives ──

def extract_narrative_facts():
    sources = []

    # Paragraph 5: Kunjam origin
    sources.append({
        "text": "The Kunjams brought all the pen from Daler and then distributed them to all the different clans at Nandraj mountain.",
        "facts": [
            {"type": "origin", "subject": "kunjam", "object": None,
             "details": "Brought pens from Daler, distributed at Nandraj mountain"}
        ]
    })

    # Paragraph 11: Brothers
    bhum_iriyal_id = resolve_id("Bhum Iriyal")
    sources.append({
        "text": "Hura Mara, Bhim Iriya, Daro Modka, Bhimaraj and Godar Bhima are all brothers.",
        "facts": [
            {"type": "sibling_group", "members": [huru_mara_id, bhum_iriyal_id, daro_moitor_id, bhimaraj_id, godar_bhima_id],
             "details": "All are brothers"}
        ]
    })
    # Create sibling relationships
    brothers = [huru_mara_id, bhum_iriyal_id, daro_moitor_id, bhimaraj_id, godar_bhima_id]
    for i in range(len(brothers)):
        for j in range(i+1, len(brothers)):
            if brothers[i] in pens and brothers[j] in pens:
                add_relationship("sibling", brothers[i], brothers[j],
                                 details="brothers", source="Paragraph 11")

    # Paragraph 15: Biriya Bhima marriage
    sources.append({
        "text": "Biriya Bhima was originally from Mangnar, and is married to Mudde Dokri, the daughter of Muddaraj",
        "facts": [
            {"type": "marriage", "from": "biriya_bhima", "to": "mudde_dokri"},
            {"type": "parent", "parent": "muddaraj", "child": "mudde_dokri",
             "details": "daughter"}
        ]
    })
    add_relationship("marriage", resolve_id("Biriya Bhima"), resolve_id("Mudde Dokri"),
                     details="Biriya Bhima married to Mudde Dokri", source="Paragraph 15")
    add_relationship("parent", resolve_id("Muddaraj"), resolve_id("Mudde Dokri"),
                     details="Muddaraj is father of Mudde Dokri", source="Paragraph 15")

    # Paragraph 17: Dol Mutte story
    # Bhimraj met Dayur Mutte (Dol Mutte), caught her hand
    # Bhogam = Chote Telam origin
    # Dol Mutte gave birth to Bhogam
    sources.append({
        "text": "Dol Mutte was pregnant... gave birth to Bhogam",
        "facts": [
            {"type": "parent", "parent": "bhimaraj", "child": "bhogam",
             "details": "Dol Mutte gave birth to Bhogam, child of Bhimaraj"},
            {"type": "relationship_note", "subject": "bhogam", "object": "telam_chote",
             "details": "Bhogams are also known as Chote Telams"}
        ]
    })
    # Bhogam is a clan, not a pen, so skip

    # Paragraph 29: Bhansi case study
    sources.append({
        "text": "Huru Mara, elder wife Urru Dokri, second wife Mawe Lungo (daughter of Nandraj and Mawli)",
        "facts": [
            {"type": "marriage", "from": "huru_mara", "to": "urru_dokri"},
            {"type": "marriage", "from": "huru_mara", "to": "mawe_lungo"},
            {"type": "parent", "parent": "nandraj", "child": "mawe_lungo"},
            {"type": "parent", "parent": "mawli", "child": "mawe_lungo"}
        ]
    })
    add_relationship("marriage", resolve_id("Huru Mara"), resolve_id("Urru Dokri"),
                     details="Huru Mara married to Urru Dokri (elder wife)", source="Paragraph 29")
    add_relationship("marriage", resolve_id("Huru Mara"), resolve_id("Mawe Lungo"),
                     details="Huru Mara married to Mawe Lungo (second wife)", source="Paragraph 29")
    add_relationship("parent", resolve_id("Nandraj"), resolve_id("Mawe Lungo"),
                     details="Nandraj is father of Mawe Lungo", source="Paragraph 29")
    add_relationship("parent", resolve_id("Mawli"), resolve_id("Mawe Lungo"),
                     details="Mawli is mother of Mawe Lungo", source="Paragraph 29")

    # Paragraph 31: Children
    sources.append({
        "text": "Punga Rungi is daughter of Huru Mara and Urru Dokri; Lug Unga is son of Huru Mara and Mawe Lungo",
        "facts": [
            {"type": "parent", "parent": "huru_mara", "child": "punga_rungi"},
            {"type": "parent", "parent": "urru_dokri", "child": "punga_rungi"},
            {"type": "parent", "parent": "huru_mara", "child": "lug_unga"},
            {"type": "parent", "parent": "mawe_lungo", "child": "lug_unga"},
        ]
    })
    for parent in ["Huru Mara", "Urru Dokri"]:
        add_relationship("parent", resolve_id(parent), resolve_id("Punga Rungi"),
                         details=f"{parent} is parent of Punga Rungi", source="Paragraph 31")
    for parent in ["Huru Mara", "Mawe Lungo"]:
        add_relationship("parent", resolve_id(parent), resolve_id("Lug Unga"),
                         details=f"{parent} is parent of Lug Unga", source="Paragraph 31")

    # Paragraph 33: Mai Sunga ghar jamai
    sources.append({
        "text": "Mai Sunga married to Punga Rungi, ghar jamai in Bhansi",
        "facts": [
            {"type": "ghar_jamai", "from": "mai_sunga", "to": "punga_rungi"},
            {"type": "marriage", "from": "mai_sunga", "to": "punga_rungi"}
        ]
    })
    add_relationship("ghar_jamai", resolve_id("Mai Sunga"), resolve_id("Punga Rungi"),
                     details="Mai Sunga is ghar jamai to Punga Rungi", source="Paragraph 33")
    add_relationship("marriage", resolve_id("Mai Sunga"), resolve_id("Punga Rungi"),
                     details="Mai Sunga married to Punga Rungi", source="Paragraph 33")

    # Paragraph 44: Godar Hunga
    sources.append({
        "text": "Godar Hunga, s/o Huru Mara and Katta Bodke (a third wife)",
        "facts": [
            {"type": "parent", "parent": "huru_mara", "child": "godar_hunga"},
            {"type": "parent", "parent": "katta_bodke", "child": "godar_hunga"},
        ]
    })
    add_relationship("parent", resolve_id("Huru Mara"), resolve_id("Godar Hunga"),
                     details="Godar Hunga is son of Huru Mara", source="Paragraph 44")
    add_relationship("parent", resolve_id("Katta Bodke"), resolve_id("Godar Hunga"),
                     details="Godar Hunga is son of Katta Bodke", source="Paragraph 44")

    # Paragraph 53: Madvi 3 brothers
    sources.append({
        "text": "3 brothers: Andal Kosa (Madvi), Gaddi Kama (Oyam), Madkaraj (Barse)",
        "facts": [
            {"type": "sibling_group", "members": ["andalkosa", "gaddi_kama", "markaraj"],
             "details": "Three brothers: Madvi, Oyam, Barse"}
        ]
    })
    madvi_bros = ["andalkosa", "gaddi_kama", "markaraj"]
    for i in range(len(madvi_bros)):
        for j in range(i+1, len(madvi_bros)):
            if madvi_bros[i] in pens and madvi_bros[j] in pens:
                add_relationship("sibling", madvi_bros[i], madvi_bros[j],
                                 details="brothers (Madvi Phratry)", source="Paragraph 53")

    # Paragraph 59: Gaddi Kama siblings
    sources.append({
        "text": "Bros 2 & 3 (Bade Vidi and Vidi Iriyal) born in Pen Kokhra; Nos 4-6 born near Gangalur",
        "facts": []
    })
    add_relationship("sibling", resolve_id("Gaddi Kama"), resolve_id("Bade Vidi Iriyal"),
                     details="Bade Vidi Iriyal is 2nd brother of Gaddi Kama (GK)", source="Paragraph 59")
    add_relationship("sibling", resolve_id("Gaddi Kama"), resolve_id("Vidi Iriyal Panda Hadma"),
                     details="Vidi Iriyal Panda Hadma is 3rd brother (youngest unmarried)", source="Table 2 Row 3")

    # Paragraph 68: Markaraj marriage
    sources.append({
        "text": "Markaraj married to Hinge Dokri. Lingal Denga is brother. Hingal Dokri represented by iron ring.",
        "facts": [
            {"type": "marriage", "from": "markaraj", "to": "hinge_dokri"}
        ]
    })
    add_relationship("marriage", resolve_id("Markaraj"), resolve_id("Hinge Dokri"),
                     details="Markaraj married to Hinge Dokri", source="Paragraph 68")
    add_relationship("sibling", resolve_id("Markaraj"), resolve_id("Lingal Denga"),
                     details="Lingal Denga is brother of Markaraj", source="Paragraph 68")

    # Paragraph 70: Ural Gunda
    sources.append({
        "text": "Ural Gunda, biggest brother, left Hingal Dokri, then Markaraj brought her",
        "facts": [
            {"type": "sibling", "from": "ural_gunda", "to": "markaraj",
             "details": "Ural Gunda is biggest brother"},
        ]
    })
    add_relationship("sibling", resolve_id("Ural Gunda"), resolve_id("Markaraj"),
                     details="Ural Gunda is biggest brother", source="Paragraph 70")
    add_relationship("marriage", resolve_id("Ural Gunda"), resolve_id("Hinge Dokri"),
                     details="Ural Gunda was earlier married to Hingal Dokri but left her",
                     source="Paragraph 70")

    # Paragraph 70: Hingal Dokri is daughter of Huru Mara
    add_relationship("parent", resolve_id("Huru Mara"), resolve_id("Hinge Dokri"),
                     details="Hinge Dokri is the daughter of Huru Mara of Bhansi",
                     source="Paragraph 70")

    # Bhime is daughter of Gadi Kama (Table 1 Row 7)
    add_relationship("parent", resolve_id("Gaddi Kama"), resolve_id("Bhime"),
                     details="Bhime is daughter of Gadi Kama (Oyam)", source="Table 1 row 7")
    add_relationship("marriage", resolve_id("Godar Bhima"), resolve_id("Bhime"),
                     details="Godar Bhima married to Bhime", source="Table 1 row 7")

    # Inge Dokri is daughter of Godar Bhima (Table 1 row 7)
    add_relationship("parent", resolve_id("Godar Bhima"), resolve_id("Inge Dokri"),
                     details="Inge Dokri is daughter of Godar Bhima, has her pen in Kuper",
                     source="Table 1 row 7")

    # Kohli Dokri is daughter of Bhum Iriyal (Table 1 row 21)
    add_relationship("parent", resolve_id("Bhum Iriyal"), resolve_id("Kohli Dokri"),
                     details="Kohli Dokri is daughter of Bhum Iriyal", source="Table 1 row 21")
    add_relationship("marriage", resolve_id("Gaddi Kama"), resolve_id("Kohli Dokri"),
                     details="Kohli Dokri married to Gaddi Kama", source="Table 2 row 4")

    # Murke Nango is daughter of Huru Mara (Table 2 row 4)
    add_relationship("parent", resolve_id("Huru Mara"), resolve_id("Murke Nango"),
                     details="Murke Nango is daughter of Huru Mara of Bhansi",
                     source="Table 2 row 4")
    add_relationship("marriage", resolve_id("Gaddi Kama"), resolve_id("Murke Nango"),
                     details="Gaddi Kama married to Murke Nango", source="Table 2 row 4")

    # Gaddi Kama's children
    for child in ["Gal Dullo", "Gal Bomda"]:
        add_relationship("parent", resolve_id("Gaddi Kama"), resolve_id(child),
                         details=f"{child} is child of Gaddi Kama and Murke Nango", source="Table 2 row 4")

    # Bomul Ungal is s/o Biriya Bhima (Table 1 row 6)
    add_relationship("parent", resolve_id("Biriya Bhima"), resolve_id("Bomul Ungal"),
                     details="Bomul Ungal is s/o Biriya Bhima", source="Table 1 row 6")

    # Kohla Kosu is s/o Raibandal (Table 1 row 17)
    add_relationship("parent", resolve_id("Raibandal"), resolve_id("Kohla Kosu"),
                     details="Kohla Kosu is s/o Raibandal", source="Table 1 row 17")

    # Vange Dokri married to Nanga Bhima (Table 3 row 1)
    add_relationship("marriage", resolve_id("Vange Dokri"), resolve_id("Nanga Bhima"),
                     details="Vange Dokri married to Nanga Bhima in Gangalur",
                     source="Table 3 row 1")

    # Hunga Moitor's wives (Table 1 row 2)
    for wife in ["Gujje Dokri", "Murde Moyo"]:
        sid = resolve_id(wife)
        if sid:
            add_relationship("marriage", resolve_id("Hunga Moitor"), sid,
                             details=f"Hunga Moitor's wife: {wife}", source="Table 1 row 2")

    # Bhimaraj married to Urru Ponde and Dol Mutte (Table 1 row 4)
    for wife in ["Urru Ponde", "Dol Mutte"]:
        add_relationship("marriage", resolve_id("Bhimaraj"), resolve_id(wife),
                         details=f"Bhimaraj married to {wife}", source="Table 1 row 4")

    return sources

narrative_sources = extract_narrative_facts()

# ── 5. Post-process known data quality issues ──

# Fix known palli values from source document that the regex couldn't parse
KNOWN_PALLI = {
    "huru_mara": ["Jhirka", "Dokometta", "Dhanora", "Tumnar", "Bhansi"],
    "daro_moitor": ["Chandenar", "Renganar"],
}
for pid, palli_list in KNOWN_PALLI.items():
    if pid in pens:
        pens[pid]["palli"] = palli_list

# Fix muddaraj's palli — move prose note to notes
if "muddaraj" in pens:
    p = pens["muddaraj"]
    cleaned = [item for item in p.get("palli", []) if len(item) <= 20]
    notes_added = [item for item in p.get("palli", []) if len(item) > 20]
    p["palli"] = cleaned if cleaned else ["Madpal"]
    for note in notes_added:
        if note not in p.get("notes", ""):
            p["notes"] = (p["notes"] + "; " if p["notes"] else "") + note

# Fix gaddi_kama karsad (mentioned in source docx)
if "gaddi_kama" in pens:
    if not pens["gaddi_kama"].get("karsad"):
        pens["gaddi_kama"]["karsad"] = "Tuesday in February"

# Clean up stale "s" entries in any remaining palli arrays
for pid, p in pens.items():
    p["palli"] = [v for v in p.get("palli", []) if v.strip() not in ("s", "S", "")]

# ── 5b. Deduplicate villages ──
# Merge entries with identical (name, phratry_id, clan_id, main_pen_id)
# by combining notes and subordinate_pen_ids into the first occurrence.
seen_village_keys = {}
unique_villages = []
for v in villages:
    key = (v["name"], v["phratry_id"], v["clan_id"], v.get("main_pen_id"))
    if key in seen_village_keys:
        existing = unique_villages[seen_village_keys[key]]
        if v.get("notes") and v["notes"] not in existing.get("notes", ""):
            existing["notes"] = (existing["notes"] + "; " if existing["notes"] else "") + v["notes"]
        for spid in v.get("subordinate_pen_ids", []):
            if spid not in existing["subordinate_pen_ids"]:
                existing["subordinate_pen_ids"].append(spid)
    else:
        seen_village_keys[key] = len(unique_villages)
        unique_villages.append(v)
villages = unique_villages

# ── 6. Build final output ──
output = {
    "metadata": {
        "title": "List of Clan Gods and Villages in Dantewada",
        "source_file": "List of Clan Gods and Villages.docx",
        "extraction_date": "2026-06-29",
        "phratry_colors": {
            "kuhrami_kadiari": {"fill": "#e74c3c", "label": "Kuhrami/Kadiari"},
            "markami": {"fill": "#3498db", "label": "Markami Kutumb"},
            "madvi": {"fill": "#2ecc71", "label": "Madvi"},
            "kawasi": {"fill": "#f39c12", "label": "Kawasi"}
        }
    },
    "phratries": list(phratries.values()),
    "clans": list(clans.values()),
    "pens": list(pens.values()),
    "villages": villages,
    "relationships": relationships,
    "narrative_sources": narrative_sources
}

# Clean up any duplicate relationships
seen_rels = set()
unique_rels = []
for rel in output["relationships"]:
    key = (rel["type"], rel["from_pen_id"], rel["to_pen_id"], rel["details"][:50])
    if key not in seen_rels:
        seen_rels.add(key)
        unique_rels.append(rel)
output["relationships"] = unique_rels

# Write output
with open(OUT_PATH, 'w', encoding='utf-8') as f:
    json.dump(output, f, indent=2, ensure_ascii=False)

print(f"✓ Wrote {OUT_PATH}")
print(f"  Phratries: {len(output['phratries'])}")
print(f"  Clans: {len(output['clans'])}")
print(f"  Pens: {len(output['pens'])}")
print(f"  Villages: {len(output['villages'])}")
print(f"  Relationships: {len(output['relationships'])}")
